import numpy as np
import docx

# ==========================================
# 1. Extração de Dados
# ==========================================
# Lê as tabelas diretamente do arquivo docx
doc = docx.Document(r'c:\Users\Breno\LAB_IA\RBF1\context\RBF1.docx')
table_test = doc.tables[4]
table_train = doc.tables[5]

def extract_table(table):
    data = []
    # Pula o cabeçalho
    for row in table.rows[1:]:
        row_data = [cell.text.strip() for cell in row.cells]
        if row_data[0] == '' or 'Taxa' in row_data[0]:
            continue
        # x1, x2, d
        data.append([float(row_data[1]), float(row_data[2]), float(row_data[3])])
    return np.array(data)

train_data = extract_table(table_train)
test_data = extract_table(table_test)

X_train = train_data[:, :2]
d_train = train_data[:, 2]

X_test = test_data[:, :2]
d_test = test_data[:, 2]

# ==========================================
# 2. Treinamento da Camada Escondida (K-means)
# ==========================================
# Filtrar apenas os padrões com presença de radiação (d = 1)
X_rad = X_train[d_train == 1]

# K-means customizado para 2 clusters
np.random.seed(42)
# Inicializar centros aleatoriamente a partir dos pontos
indices = np.random.choice(X_rad.shape[0], 2, replace=False)
centers = X_rad[indices].copy()
labels = np.zeros(X_rad.shape[0])

for _ in range(100): # max_iter
    # Atribuir rótulos
    dists = np.zeros((X_rad.shape[0], 2))
    for i in range(2):
        dists[:, i] = np.sum((X_rad - centers[i])**2, axis=1)
    new_labels = np.argmin(dists, axis=1)
    
    if np.all(labels == new_labels):
        break
    labels = new_labels
    
    # Atualizar centros
    for i in range(2):
        if np.sum(labels == i) > 0:
            centers[i] = np.mean(X_rad[labels == i], axis=0)

# Calcular variâncias para cada cluster
variances = np.zeros(2)
for i in range(2):
    cluster_points = X_rad[labels == i]
    # Variância = média do quadrado das distâncias ao centro
    dists_sq = np.sum((cluster_points - centers[i])**2, axis=1)
    variances[i] = np.mean(dists_sq)
    if variances[i] == 0:
        variances[i] = 1e-6 # evitar divisão por zero

print("Centros:\n", centers)
print("Variâncias:", variances)

# ==========================================
# 3. Calcular saídas da camada escondida (RBF)
# ==========================================
def gaussian_rbf(X, centers, variances):
    phi = np.zeros((X.shape[0], len(centers)))
    for i in range(len(centers)):
        dists_sq = np.sum((X - centers[i])**2, axis=1)
        phi[:, i] = np.exp(-dists_sq / (2 * variances[i]))
    return phi

Phi_train = gaussian_rbf(X_train, centers, variances)
# Adicionar bias (-1)
Phi_train_bias = np.hstack((-np.ones((Phi_train.shape[0], 1)), Phi_train))

# ==========================================
# 4. Treinamento da Camada de Saída (Regra Delta)
# ==========================================
eta = 0.01
epsilon = 1e-7

np.random.seed(42)
w = np.random.rand(3) # [W21,0 (bias), W21,1, W21,2]

epoch = 0
prev_mse = float('inf')

while True:
    mse = 0
    # Calcular o MSE da época atual
    for i in range(Phi_train_bias.shape[0]):
        phi = Phi_train_bias[i]
        d = d_train[i]
        v = np.dot(w, phi)
        mse += (d - v)**2
    mse /= Phi_train_bias.shape[0]
    
    # Critério de parada
    if abs(prev_mse - mse) < epsilon:
        break
    prev_mse = mse
    
    # Atualização de pesos
    for i in range(Phi_train_bias.shape[0]):
        phi = Phi_train_bias[i]
        d = d_train[i]
        v = np.dot(w, phi)
        e = d - v
        w = w + eta * e * phi
        
    epoch += 1

print(f"Treinamento concluído em {epoch} épocas.")
print(f"Pesos finais: W21,0={w[0]:.4f}, W21,1={w[1]:.4f}, W21,2={w[2]:.4f}")

# ==========================================
# 5. Validação (Conjunto de Teste)
# ==========================================
Phi_test = gaussian_rbf(X_test, centers, variances)
Phi_test_bias = np.hstack((-np.ones((Phi_test.shape[0], 1)), Phi_test))

predictions = []
predictions_real = []
for i in range(Phi_test_bias.shape[0]):
    phi = Phi_test_bias[i]
    v = np.dot(w, phi)
    predictions_real.append(v)
    y_pos = 1 if v >= 0 else -1
    predictions.append(y_pos)

predictions = np.array(predictions)
accuracy = np.mean(predictions == d_test) * 100

print(f"Taxa de acerto: {accuracy:.2f}%")
