import docx
import numpy as np
import matplotlib.pyplot as plt
import os

# Set working directory to script location
os.chdir(os.path.dirname(os.path.abspath(__file__)))

# ==========================================
# 1. Extração de Dados
# ==========================================
doc = docx.Document(r'c:\Users\Breno\LAB_IA\RBF2\context\RBF2.docx')
table_test = doc.tables[2]
table_train = doc.tables[3]

train_data = []
# Skip header
for row in table_train.rows[1:]:
    cells = [c.text.strip().replace(',', '.') for c in row.cells]
    for i in range(3):
        idx = i*5
        if cells[idx] == '': continue
        train_data.append([float(cells[idx+1]), float(cells[idx+2]), float(cells[idx+3]), float(cells[idx+4])])

test_data = []
# Rows 2 to 16 (15 samples)
for row in table_test.rows[2:17]:
    cells = [c.text.strip().replace(',', '.') for c in row.cells]
    test_data.append([float(cells[1]), float(cells[2]), float(cells[3]), float(cells[4])])

train_data = np.array(train_data)
test_data = np.array(test_data)

X_train = train_data[:, :3]
d_train = train_data[:, 3]

X_test = test_data[:, :3]
d_test = test_data[:, 3]

# ==========================================
# Funções Auxiliares
# ==========================================
def kmeans(X, k, seed):
    np.random.seed(seed)
    indices = np.random.choice(X.shape[0], k, replace=False)
    centers = X[indices].copy()
    labels = np.zeros(X.shape[0])
    
    for _ in range(200):
        dists = np.zeros((X.shape[0], k))
        for i in range(k):
            dists[:, i] = np.sum((X - centers[i])**2, axis=1)
        new_labels = np.argmin(dists, axis=1)
        
        if np.all(labels == new_labels):
            break
        labels = new_labels
        
        for i in range(k):
            if np.sum(labels == i) > 0:
                centers[i] = np.mean(X[labels == i], axis=0)
                
    variances = np.zeros(k)
    for i in range(k):
        cluster_points = X[labels == i]
        if len(cluster_points) > 0:
            dists_sq = np.sum((cluster_points - centers[i])**2, axis=1)
            variances[i] = np.mean(dists_sq)
        if variances[i] == 0:
            variances[i] = 1e-6
            
    return centers, variances

def gaussian_rbf(X, centers, variances):
    phi = np.zeros((X.shape[0], len(centers)))
    for i in range(len(centers)):
        dists_sq = np.sum((X - centers[i])**2, axis=1)
        phi[:, i] = np.exp(-dists_sq / (2 * variances[i]))
    return phi

# ==========================================
# Execução
# ==========================================
topologies = [5, 10, 15]
results = {}

plt.figure(figsize=(15, 5))

for idx_top, N1 in enumerate(topologies):
    print(f"\n--- Rede {idx_top+1} (N1={N1}) ---")
    results[N1] = []
    
    best_mse = float('inf')
    best_mse_history = []
    best_t = -1
    
    for t in range(3):
        seed = 42 + t*10 + N1
        
        # 1. K-Means
        centers, variances = kmeans(X_train, N1, seed)
        
        # 2. RBF
        Phi_train = gaussian_rbf(X_train, centers, variances)
        Phi_train_bias = np.hstack((-np.ones((Phi_train.shape[0], 1)), Phi_train))
        
        Phi_test = gaussian_rbf(X_test, centers, variances)
        Phi_test_bias = np.hstack((-np.ones((Phi_test.shape[0], 1)), Phi_test))
        
        # 3. Adaline / Delta Rule
        eta = 0.01
        epsilon = 1e-7
        np.random.seed(seed)
        w = np.random.rand(N1 + 1) # values between 0 and 1
        
        epoch = 0
        prev_mse = float('inf')
        mse_history = []
        
        while True:
            # Predictions
            v = np.dot(Phi_train_bias, w)
            e = d_train - v
            mse = np.mean(e**2)
            mse_history.append(mse)
            
            if abs(prev_mse - mse) < epsilon:
                break
            prev_mse = mse
            
            # Batch update (or stochastic? standard delta rule is batch or stochastic. Adaline is usually batch or stochastic. Let's do stochastic since it's common, or batch. "A regra delta generalizada" often means stochastic. Wait, in RBF1 we did stochastic. Let's do stochastic to match RBF1).
            for i in range(Phi_train_bias.shape[0]):
                v_i = np.dot(w, Phi_train_bias[i])
                e_i = d_train[i] - v_i
                w += eta * e_i * Phi_train_bias[i]
                
            epoch += 1
            if epoch > 5000: # safety break
                break
                
        # Calculate Final Train MSE using final weights
        v = np.dot(Phi_train_bias, w)
        final_mse = np.mean((d_train - v)**2)
        
        # Validate
        v_test = np.dot(Phi_test_bias, w)
        relative_errors = np.abs(v_test - d_test) / np.abs(d_test) * 100
        mre = np.mean(relative_errors)
        var_mre = np.var(relative_errors) # variance in % squared. usually just standard variance of the relative errors.
        
        print(f" Treinamento {t+1}: Épocas={epoch}, EQM={final_mse:.6f}, MRE={mre:.2f}%, Var={var_mre:.2f}%")
        
        results[N1].append({
            'epochs': epoch,
            'mse': final_mse,
            'mre': mre,
            'var_mre': var_mre,
            'y_test': v_test,
            'mse_history': mse_history
        })
        
        # Track best
        if final_mse < best_mse:
            best_mse = final_mse
            best_mse_history = mse_history
            best_t = t
            
    # Plot best
    plt.subplot(1, 3, idx_top+1)
    plt.plot(best_mse_history, color='blue')
    plt.title(f'Rede {idx_top+1} (N1={N1})\nMelhor Treinamento: T{best_t+1}')
    plt.xlabel('Épocas')
    plt.ylabel('EQM')
    plt.grid(True)

plt.tight_layout()
plt.savefig('graficos_mse.png')
print("\nGráficos salvos em graficos_mse.png")

# Print output formatting for markdown
print("\n--- RESULTADOS PARA MARKDOWN ---")
for idx, N1 in enumerate(topologies):
    print(f"Rede {idx+1} (N1={N1})")
    for t in range(3):
        res = results[N1][t]
        print(f"T{t+1}: EQM={res['mse']:.6f}, Epocas={res['epochs']}")
        
for i in range(15):
    row_str = f"| {i+1} | {X_test[i][0]} | {X_test[i][1]} | {X_test[i][2]} | {d_test[i]} |"
    for N1 in topologies:
        for t in range(3):
            y_pred = results[N1][t]['y_test'][i]
            row_str += f" {y_pred:.4f} |"
    print(row_str)

print("MRE e Var:")
for N1 in topologies:
    print(f"N1={N1}:")
    for t in range(3):
        res = results[N1][t]
        print(f"  T{t+1}: MRE={res['mre']:.2f}%, Var={res['var_mre']:.2f}%")
